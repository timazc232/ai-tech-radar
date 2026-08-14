# -*- coding: utf-8 -*-
"""Convert the V1.1 design markdown (headings / tables / bullets / paragraphs) to .docx."""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn


def set_east_asian(doc):
    names = ['Normal', 'Title', 'Heading 1', 'Heading 2', 'Heading 3', 'List Bullet']
    for n in names:
        try:
            st = doc.styles[n]
        except KeyError:
            continue
        st.font.name = 'Calibri'
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.append(rfonts)
        rfonts.set(qn('w:eastAsia'), '微软雅黑')
        if n == 'Normal':
            st.font.size = Pt(10.5)
        if n.startswith('Heading'):
            st.font.color.rgb = RGBColor(0x1F, 0x3B, 0x63)


def strip_md(text):
    return text.replace('**', '').strip()


def add_table(doc, rows):
    cells = []
    for line in rows:
        parts = [strip_md(c) for c in line.strip().strip('|').split('|')]
        cells.append(parts)
    ncols = max(len(r) for r in cells)
    table = doc.add_table(rows=len(cells), cols=ncols)
    table.style = 'Table Grid'
    for i, row in enumerate(cells):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.text = row[j] if j < len(row) else ''
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    if i == 0:
                        r.bold = True
    doc.add_paragraph()


def add_code_block(doc, code_lines):
    """Render a fenced code block as a monospace, shaded paragraph with real line breaks."""
    p = doc.add_paragraph()
    def add_run(text):
        r = p.add_run(text)
        r.font.name = 'Consolas'
        rpr = r._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.append(rfonts)
        rfonts.set(qn('w:ascii'), 'Consolas')
        rfonts.set(qn('w:hAnsi'), 'Consolas')
        rfonts.set(qn('w:eastAsia'), 'Consolas')
        r.font.size = Pt(9)
        return r
    for idx, ln in enumerate(code_lines):
        if idx > 0:
            add_run('').add_break()   # <w:br/> -> real line break in Word
        add_run(ln if ln != '' else ' ')
    # light gray shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {})
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)


def convert(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()
    set_east_asian(doc)
    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    i = 0
    seen_heading = False
    table_buf = []
    code_buf = []
    in_code = False

    def flush_table():
        if table_buf:
            add_table(doc, table_buf[:])
            table_buf.clear()

    def flush_code():
        if code_buf:
            add_code_block(doc, code_buf[:])
            code_buf.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code block toggle
        if stripped.startswith('```'):
            if not in_code:
                flush_table()
                in_code = True
            else:
                flush_code()
                in_code = False
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped.startswith('|'):
            table_buf.append(stripped)
            i += 1
            continue
        flush_table()

        if not stripped:
            i += 1
            continue

        if stripped.startswith('#'):
            seen_heading = True
            level = len(stripped) - len(stripped.lstrip('#'))
            level = min(level, 3)
            doc.add_heading(strip_md(stripped.lstrip('#')), level=level)
        elif stripped.startswith('- '):
            doc.add_paragraph(strip_md(stripped[2:]), style='List Bullet')
        elif not seen_heading:
            # front-matter lines before first heading -> title block
            p = doc.add_paragraph()
            r = p.add_run(strip_md(stripped))
            if stripped.startswith('PRODUCT'):
                r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(0x5A, 0x6B, 0x84)
            elif stripped.startswith('AI Tech Radar'):
                r.font.size = Pt(26)
                r.bold = True
            else:
                r.font.size = Pt(14)
                r.bold = True
        else:
            doc.add_paragraph(strip_md(stripped))
        i += 1

    flush_table()
    flush_code()
    doc.save(docx_path)
    print('saved:', docx_path)


if __name__ == '__main__':
    convert(sys.argv[1], sys.argv[2])
